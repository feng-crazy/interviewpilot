from io import BytesIO
import re
import pdfplumber
from docx import Document

from ..config.logging import get_logger


class ResumeParserService:
    """Service for extracting text from PDF and DOCX resume files."""

    def __init__(self):
        self.logger = get_logger("resume_parser")

    async def extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from PDF or DOCX resume."""
        ext = filename.lower().split(".")[-1]

        if ext == "pdf":
            return self._extract_pdf(content)
        elif ext in ["docx", "doc"]:
            return self._extract_docx(content)
        else:
            raise ValueError(f"Unsupported format: {ext}")

    def _extract_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF using pdfplumber with character-level filtering.

        Some PDFs contain hidden watermarks or anti-copy characters that disrupt
        layout-based text extraction. This method extracts characters directly
        and filters out suspicious fonts/characters.
        """
        self.logger.info("resume_parse_start", file_type="pdf")
        text_parts = []
        try:
            with pdfplumber.open(BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(
                        layout=True,
                        x_tolerance=3,
                        y_tolerance=3,
                    )

                    if page_text and self._is_text_corrupted(page_text):
                        self.logger.info(
                            "pdf_layout_extraction_corrupted",
                            fallback="character_extraction",
                        )
                        page_text = self._extract_chars_filtered(page)

                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {str(e)}")

        text = "\n\n".join(text_parts)
        text = self._remove_watermark_suffix(text)
        self.logger.info("resume_parse_end", file_type="pdf", text_length=len(text))
        return text

    def _remove_watermark_suffix(self, text: str) -> str:
        """
        Remove watermark/anti-copy suffix from extracted text.

        PDF watermarks often appear at the end of extracted text as patterns like:
        - ~~~ followed by encoded strings
        - Strings containing ~, _, and hex-like characters mixed together
        - Examples: "~~A3rhhNQT_liGOWXy_U_WoyYNlF7m934dH1e31603ed6b764fdb"

        This method detects and removes such watermark suffixes.
        """
        if not text:
            return text

        # Pattern 1: Remove trailing watermark strings (~~ + encoded pattern)
        # Matches: ~~, followed by any combination of ~, _, letters, digits, up to end
        watermark_pattern = re.compile(r"[~_]{2,}[a-zA-Z0-9~_]{10,}$")

        cleaned = watermark_pattern.sub("", text)

        # Pattern 2: Remove standalone watermark blocks at the very end
        # These are typically hex-like strings with ~ and _ interspersed
        # Example: "bdf467b6de30613e1Hd439m7FlNYyoW_U_yXWOGil_TQNhhr3A"
        standalone_watermark = re.compile(r"[a-f0-9A-Z~_]{20,}$")

        cleaned = standalone_watermark.sub("", cleaned)

        # Pattern 3: Remove repeated watermark patterns
        # Sometimes watermarks repeat multiple times at the end
        repeated_pattern = re.compile(r"(?:[a-zA-Z0-9~_]{15,}){2,}$")

        cleaned = repeated_pattern.sub("", cleaned)

        # Trim trailing whitespace and newlines
        cleaned = cleaned.rstrip()

        return cleaned

    def _is_text_corrupted(self, text: str) -> bool:
        """
        Detect if extracted text appears corrupted by watermarks/anti-copy.

        Patterns indicating corruption:
        - High density of '~' and '_' characters in first region
        - Random hex-like strings (e.g., bdf467b6de306...)
        - Leading content after whitespace is '~' or '_'
        """
        if not text:
            return False

        first_200 = text[:200]

        weird_count = first_200.count("~") + first_200.count("_")
        if weird_count > 3:
            return True

        if re.search(r"[a-f0-9]{12,}", first_200, re.IGNORECASE):
            return True

        stripped = first_200.strip()
        if stripped.startswith("~") or stripped.startswith("_"):
            return True

        return False

    def _extract_chars_filtered(self, page) -> str:
        """
        Extract text by filtering individual characters.

        Filters out characters from suspicious fonts (Helvetica, Arial, Times)
        that match watermark patterns (~, _, hex strings).
        """
        chars = page.chars
        if not chars:
            return ""

        valid_pattern = re.compile(
            r"[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]"
            r"|[a-zA-Z0-9]"
            r"|[\s\|、：；，。！？\-—()（）【】\[\]·•/]"
        )

        suspicious_fonts = ["Helvetica", "Arial", "Times"]

        filtered_chars = []
        for c in chars:
            text = c.get("text", "")
            fontname = c.get("fontname", "")

            if not text.strip():
                filtered_chars.append(text)
                continue

            if any(sf in fontname for sf in suspicious_fonts):
                if text in ["~", "_"] or re.match(
                    r"^[a-f0-9]{6,}$", text, re.IGNORECASE
                ):
                    continue

            if valid_pattern.match(text):
                filtered_chars.append(text)

        return "".join(filtered_chars)

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        self.logger.info("resume_parse_start", file_type="docx")
        text_parts = []
        try:
            doc = Document(BytesIO(content))

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {str(e)}")

        text = "\n\n".join(text_parts)
        self.logger.info("resume_parse_end", file_type="docx", text_length=len(text))
        return text
